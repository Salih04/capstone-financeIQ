from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1A, 0x3A, 0x5C)   # headings
MID_BLUE   = RGBColor(0x27, 0x6F, 0xBF)   # sub-headings / table headers
LIGHT_BLUE = RGBColor(0xD6, 0xE8, 0xF7)   # table header fill
ALT_ROW    = RGBColor(0xF0, 0xF6, 0xFD)   # alternating row fill
GREEN      = RGBColor(0x1E, 0x7E, 0x34)
ORANGE     = RGBColor(0xCC, 0x70, 0x00)
GRAY_TEXT  = RGBColor(0x55, 0x55, 0x55)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helpers ───────────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="C0C0C0"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top","left","bottom","right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    # top border decoration
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "276FBF")
    pBdr.append(bottom)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_BLUE
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = MID_BLUE
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_BLUE
    return p

def body(text, italic=False, color=None):
    p   = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size   = Pt(10)
    run.italic      = italic
    if color:
        run.font.color.rgb = color
    return p

def bullet(text, level=0):
    p   = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def code_block(text):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.3)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    return p

def make_table(headers, rows, col_widths=None):
    tbl  = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header row
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        set_cell_bg(cell, MID_BLUE)
        set_cell_borders(cell, "276FBF")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p   = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold           = True
        run.font.size      = Pt(9)
        run.font.color.rgb = WHITE

    # data rows
    for ri, row_data in enumerate(rows):
        cells = tbl.rows[ri + 1].cells
        bg    = ALT_ROW if ri % 2 == 1 else WHITE
        for ci, val in enumerate(row_data):
            cell = cells[ci]
            set_cell_bg(cell, bg)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p   = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)

    # column widths
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacing after table
    return tbl

def page_break():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run("FinanceIQ")
run.bold           = True
run.font.size      = Pt(32)
run.font.color.rgb = DARK_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Capstone Project — Technical Progress Report")
run.bold           = True
run.font.size      = Pt(16)
run.font.color.rgb = MID_BLUE

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("BIST Stock Scoring & Financial Analytics Platform")
run.font.size      = Pt(12)
run.italic         = True
run.font.color.rgb = GRAY_TEXT

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"Report Date: {datetime.date.today().strftime('%B %d, %Y')}")
run.font.size      = Pt(11)
run.font.color.rgb = GRAY_TEXT

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Repository: salih04/capstone-financeiq")
run.font.size      = Pt(11)
run.font.color.rgb = GRAY_TEXT

page_break()

# ══════════════════════════════════════════════════════════════
#  SECTION 1 — SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════
heading1("1. System Architecture")

heading2("1.1 Frontend → Backend → Database Flow")
code_block(
    "User Browser\n"
    "    │\n"
    "    ▼\n"
    "React SPA  (port 3000)\n"
    "    │   Vite build · React 18 · React Router 6\n"
    "    │   axios client with automatic JWT Bearer injection\n"
    "    │   /api/* requests proxied via Nginx\n"
    "    │\n"
    "    ▼\n"
    "Nginx Reverse Proxy  (port 80 inside container)\n"
    "    │   proxy_pass → backend:8000\n"
    "    │   SPA fallback: all routes → index.html\n"
    "    │\n"
    "    ▼\n"
    "FastAPI Backend  (port 8000)\n"
    "    │   11 routers · 50+ endpoints · 17 service modules\n"
    "    │   SQLAlchemy ORM sessions · Pydantic v2 validation\n"
    "    │\n"
    "    ▼\n"
    "PostgreSQL 16  (port 5432)\n"
    "    25+ tables: users, companies, financial_statements,\n"
    "    computed_metrics, score_runs, forecast_runs, audit_logs …"
)

heading2("1.2 API Structure")
body("11 routers are registered on the FastAPI application:")
make_table(
    ["Router", "Prefix / Mount", "Key Operations"],
    [
        ["Auth",          "/auth",        "register, login, me (JWT)"],
        ["Users",         "/users",       "profile read / update, onboarding fields"],
        ["Companies",     "/companies",   "search, financials, metrics, transitions, sector-scores"],
        ["Scoring",       "/",            "single score, compare, adaptive weights, score-run retrieval"],
        ["Financials",    "/financials",  "CSV import, 12-ratio auto-calculation"],
        ["Forecasting",   "/",            "upload presets, train, rank predictions, rolling CV evaluate"],
        ["Reports",       "/reports",     "export CSV / JSON / PDF per score-run"],
        ["Ingestion",     "/ingestion",   "V3 pipeline, job history, data-health dashboard"],
        ["Validation",    "/validation",  "time-split validation, feature importances (SHAP)"],
        ["Labeling",      "/labeling",    "label definition CRUD, distribution preview, activate"],
        ["Admin",         "/admin",       "model registry, audit logs, user management"],
        ["Fundamentals",  "/",            "quarterly fundamentals CSV, KAP ratios"],
        ["News",          "/news",        "Finnhub market news (12 h cache)"],
        ["Health",        "/health",      "API status + version"],
    ],
    col_widths=[1.3, 1.4, 3.5]
)

heading2("1.3 Authentication Structure")
code_block(
    "POST /auth/login\n"
    "    ├── Lookup user by email (PostgreSQL)\n"
    "    ├── Check account lockout (5 failed attempts → 15-min lock)\n"
    "    ├── Verify bcrypt password hash  (passlib, cost = 12)\n"
    "    ├── Issue JWT  (HS256, 24 h expiry)  via python-jose\n"
    "    └── Reset failed_login_count on success\n"
    "\n"
    "Protected Endpoints\n"
    "    └── Authorization: Bearer <token>\n"
    "            └── get_current_user dependency\n"
    "                    ├── Decode & validate JWT signature\n"
    "                    ├── Fetch live user record from DB\n"
    "                    └── Enforce role  (investor | admin | analyst)\n"
    "\n"
    "Frontend\n"
    "    ├── Token stored in localStorage\n"
    "    ├── axios interceptor injects token on every request\n"
    "    └── 401 response → automatic redirect to /login"
)

heading2("1.4 Deployment Structure")
code_block(
    "docker-compose.yml\n"
    "├── db        postgres:16-alpine · port 5432 · persistent named volume\n"
    "│             health-check: pg_isready\n"
    "│\n"
    "├── backend   python:3.12-slim · port 8000\n"
    "│             bind-mount ./2.backend & ./3.Datasets (live reload)\n"
    "│             CMD: python seed.py && uvicorn app.main:app --reload\n"
    "│             env: DATABASE_URL, SECRET_KEY, ALGORITHM, TOKEN_EXPIRE\n"
    "│\n"
    "└── frontend  multi-stage build\n"
    "              Stage 1 – node:20-alpine  →  npm run build (Vite)\n"
    "              Stage 2 – nginx:alpine   →  serve /dist\n"
    "              Nginx: /api/* → backend:8000, SPA fallback\n"
    "              Exposed: port 3000 (host)\n"
    "\n"
    "Database migrations: Alembic  (5 migration files, 2026-04-06)\n"
    "Fallback:  Base.metadata.create_all() for missing tables\n"
    "Start:     docker-compose up --build\n"
    "Swagger:   http://localhost:8000/docs"
)

page_break()

# ══════════════════════════════════════════════════════════════
#  SECTION 2 — IMPLEMENTATION STATUS
# ══════════════════════════════════════════════════════════════
heading1("2. Current Implementation Status")

make_table(
    ["Feature", "Status", "Notes"],
    [
        ["Login / Authentication",      "✅ Complete", "JWT · bcrypt · account lockout (5 attempts/15 min) · role-based access"],
        ["Company Search",              "✅ Complete", "Filter by ticker, name, sector, dataset; detail page with metrics"],
        ["Financial Data Ingestion",    "✅ Complete", "V3 pipeline: job tracking · outlier detection · data-quality issues"],
        ["Ratio Calculation Engine",    "✅ Complete", "12 ratios auto-computed on import (ROA, ROE, margins, liquidity, leverage, CF)"],
        ["ML / Scoring Engine",         "✅ Complete", "Dual-mode (rule-based + logistic) · 5-model ensemble · adaptive weights · 3-level explainability"],
        ["Stock Comparison",            "✅ Complete", "Multi-stock comparison across 5 models + ensemble · common-period intersection"],
        ["Export — PDF",                "✅ Complete", "ReportLab styled report: summary · analysis · metric breakdown · AI commentary"],
        ["Export — CSV",                "✅ Complete", "UTF-8-SIG (Excel-compatible) · header · summary · commentary · metric breakdown"],
        ["Export — JSON",               "✅ Complete", "Pretty-printed · full score detail + AI commentary · Turkish char safe"],
        ["Admin Panel",                 "✅ Complete", "Model registry (CRUD) · audit log · user management · role promotion"],
        ["Database Integration",        "✅ Complete", "PostgreSQL 16 · SQLAlchemy 2.0 · Alembic migrations · 25+ tables"],
        ["API Endpoints",               "✅ Complete", "50+ endpoints · Swagger UI at /docs · Pydantic v2 validation"],
        ["Forecasting Engine",          "✅ Complete", "Upload 2020–2025 winners · train sector models · ranked predictions · rolling CV"],
        ["Validation Lab",              "✅ Complete", "Time-split val · confusion matrix · ROC-AUC · SHAP feature importances"],
        ["Labeling Lab",                "✅ Complete", "Label definition CRUD · sector benchmarks · distribution preview · activation"],
        ["Data Health Dashboard",       "✅ Complete", "Ingestion job history · per-row issue log (missing/outlier/duplicate/stale)"],
        ["News Feed",                   "✅ Complete", "Finnhub API integration · 12 h server-side cache"],
        ["User Onboarding",             "✅ Complete", "user_type · risk_level · investment_scope · sector_focus profile fields"],
        ["AI Natural Language Search",  "🔄 In Progress", "Route & page scaffold exist (/ai-search) · NL query logic pending"],
        ["Airflow Scheduling",          "🔄 In Progress", "forecasting_retrain_dag.py defined · Airflow not yet in docker-compose"],
    ],
    col_widths=[2.2, 1.3, 2.7]
)

page_break()

# ══════════════════════════════════════════════════════════════
#  SECTION 3 — MACHINE LEARNING DETAILS
# ══════════════════════════════════════════════════════════════
heading1("3. Machine Learning Details")

heading2("3.1 Models Being Tested / Used")
make_table(
    ["Model", "Type", "Usage"],
    [
        ["Rule-Based Scorer",                  "Weighted threshold rules",         "Default scoring (no training required) · produces 0–100 score"],
        ["Logistic Regression",                "Binary classification",            "Trained on computed metrics → success label · alternative scoring mode"],
        ["ElasticNet",                         "L1+L2 regularized regression",     "Forecasting ensemble member 1 · high interpretability"],
        ["Random Forest",                      "Bootstrap ensemble (bagging)",     "Forecasting ensemble member 2 · handles non-linearity"],
        ["XGBoost",                            "Gradient boosting (boosting)",     "Forecasting ensemble member 3 · early stopping"],
        ["SARIMAX",                            "Seasonal time-series AR model",    "Forecasting ensemble member 4 · captures temporal patterns"],
        ["Temporal Fusion Transformer (TFT)",  "Attention-based sequence model",   "Forecasting ensemble member 5 · multi-horizon forecasting"],
    ],
    col_widths=[2.0, 2.0, 2.2]
)
body("Multi-model ensemble initial weights: 0.2 each (equal). Weights updated from rolling cross-validation results.", italic=True, color=GRAY_TEXT)

heading2("3.2 Dataset")
bullet("Winner cohort Excel files: 6 files covering 2020–2025, BIST Turkish stocks")
bullet("Quarterly fundamentals: user-uploaded CSV, 30+ financial fields per stock/period")
bullet("KAP HTML financials: scraped from Turkish capital markets board disclosures")
bullet("Finnhub API: real-time market news (not used for model training)")
bullet("Time range: 2020–2025 (6 years annual; quarterly where available)")

heading2("3.3 Train / Test Methodology")
bullet("Time-Split Validation: data sorted by period; first 70% → train, last 30% → test (no data leakage)")
bullet("Rolling Window CV: configurable window size; per-fold metrics stored in forecast_evaluation_folds table")
bullet("Labeling Strategies: sector median benchmark · upper quartile · risk-adjusted · custom threshold rules")
bullet("Feature scaling: StandardScaler applied before logistic regression and tree-based models")
bullet("Missing value imputation: median per feature / sector / period (forecasting pipeline)")

heading2("3.4 Metrics Tracked")
make_table(
    ["Metric", "Description"],
    [
        ["Accuracy",          "Fraction of correctly classified samples"],
        ["Precision",         "TP / (TP + FP)"],
        ["Recall",            "TP / (TP + FN)"],
        ["F1 Score",          "Harmonic mean of precision and recall"],
        ["ROC-AUC",           "Area under receiver operating characteristic curve"],
        ["Confusion Matrix",  "TP / FP / TN / FN stored as JSON"],
        ["Calibration",       "Prediction confidence vs actual outcome summary"],
        ["Rank Stability",    "Mean overlap@K across rolling CV folds (forecasting)"],
    ],
    col_widths=[1.8, 4.4]
)
body("Quantitative metric values depend on labeled dataset loaded at run time. Validation infrastructure is fully operational.", italic=True, color=GRAY_TEXT)

heading2("3.5 Key Financial Features / Ratios Identified")
body("12 core computed metrics used across all scoring modes:")
make_table(
    ["Category", "Metric", "Rule-Based Weight"],
    [
        ["Profitability", "Return on Assets (ROA)",            "15 pts"],
        ["Profitability", "Return on Equity (ROE)",            "15 pts"],
        ["Profitability", "Operating Margin",                  "10 pts"],
        ["Profitability", "Net Margin",                        "5 pts"],
        ["Liquidity",     "Current Ratio",                     "10 pts"],
        ["Liquidity",     "Quick Ratio",                       "5 pts"],
        ["Liquidity",     "Cash Ratio",                        "5 pts"],
        ["Leverage",      "Debt-to-Equity",                    "10 pts"],
        ["Leverage",      "Debt-to-Assets",                    "5 pts"],
        ["Cash Flow",     "OCF-to-Debt",                       "10 pts"],
        ["Cash Flow",     "OCF-to-Assets",                     "5 pts"],
        ["Cash Flow",     "Cash Flow Margin",                  "5 pts"],
    ],
    col_widths=[1.6, 2.8, 1.8]
)
body("Category totals: Profitability 35 pts · Cash Flow 30 pts · Liquidity 20 pts · Leverage 15 pts = 100 pts total.", italic=True)

heading2("3.6 Explainability (3-Level Framework)")
heading3("Level 1 — Feature Contribution")
bullet("Per-metric: raw value · normalized value (z-score vs sector benchmark)")
bullet("Weight from model + contribution to total score")
bullet("Transition delta from previous period · sector z-score context")

heading3("Level 2 — Human-Readable Sentence (Turkish)")
bullet("Absolute level assessment: 'strong' / 'medium' / 'weak'")
bullet("Trend indicator: 'improving ✓' or 'deteriorating ✗'")
bullet("Sector context: 'above average' / 'below peers'")
bullet("Risk signal if applicable")

heading3("Level 3 — Counterfactual Mini-Insight")
bullet("'If metric X had been Y, total score would be Z'")
bullet("Constructed by perturbing one metric at a time")
bullet("Identifies high-impact improvement levers for the analyst")

page_break()

# ══════════════════════════════════════════════════════════════
#  SECTION 4 — DATA PIPELINE
# ══════════════════════════════════════════════════════════════
heading1("4. Data Pipeline Details")

heading2("4.1 Data Sources")
make_table(
    ["Source", "Format", "Ingested Into"],
    [
        ["Winner Excel Presets (2020–2025)", "6 × .xlsx files in 3.Datasets/", "winner_cohort_rows table"],
        ["Quarterly Fundamentals CSV",       "User upload via /fundamentals/upload-csv (30+ fields)", "quarterly_fundamentals table"],
        ["KAP HTML Financials",              "Scraped via kap_financials_service.py (BeautifulSoup4)", "financial_statements table"],
        ["Manual CSV Import",                "Structured upload via /ingestion/import/csv", "financial_statements + computed_metrics"],
        ["Finnhub REST API",                 "JSON market news (12 h server-side cache)", "Served directly (not persisted)"],
    ],
    col_widths=[2.3, 2.5, 1.4]
)

heading2("4.2 How Financial Data Is Retrieved — V3 Ingestion Pipeline")
code_block(
    "User uploads CSV / Excel file\n"
    "        │\n"
    "        ▼\n"
    "Ingestion Router  →  IngestionService (V3)\n"
    "        │\n"
    "        ├── 1. Validate required fields\n"
    "        │        (ticker, period, revenue, net_income, total_assets, total_equity)\n"
    "        │        Missing fields → DataQualityIssue(type=missing_field, severity=error)\n"
    "        │\n"
    "        ├── 2. Duplicate detection  (company_id, period) uniqueness check\n"
    "        │\n"
    "        ├── 3. Upsert FinancialStatement row\n"
    "        │\n"
    "        ├── 4. RatioService.compute_ratios() → 12 metrics\n"
    "        │        → Upsert ComputedMetric row\n"
    "        │\n"
    "        ├── 5. Outlier detection per metric\n"
    "        │        e.g. ROA ∈ (-2.0, 2.0), D/E ∈ (0, 50)\n"
    "        │        → DataQualityIssue(type=outlier, severity=warning)\n"
    "        │\n"
    "        └── 6. Post-processing\n"
    "                 ├── MetricTransition  (period-over-period deltas)\n"
    "                 ├── SectorBenchmark   (mean, std, percentiles per sector/period)\n"
    "                 └── SectorNormalizedFeature  (z-scores, percentile ranks)\n"
    "\n"
    "Job Tracking:  IngestionJob row (status: queued → running → success|partial|failed)\n"
    "               items_total / items_success / items_failed counted per run"
)

heading2("4.3 Preprocessing / Normalization Methods")
make_table(
    ["Method", "Details"],
    [
        ["Text Normalization",     "Unicode NFKC · strip whitespace · Turkish character handling"],
        ["Number Parsing",         "Handles Turkish decimal format (comma vs period) · currency symbols (₺, TL) · percentages"],
        ["Missing Value Imputation","Median per feature / sector / period (ML training pipeline)"],
        ["Feature Scaling",        "StandardScaler applied before logistic regression and tree-based models"],
        ["Outlier Handling",       "Configurable per-metric thresholds · flagged as warnings (not excluded) · user reviews via Data Health page"],
        ["Duplicate Detection",    "Single row enforced per (company, period) via DB unique constraint"],
    ],
    col_widths=[2.2, 4.0]
)

heading2("4.4 Sector Adjustment Method")
bullet("SectorBenchmark: for each (sector_code, period, feature_name) compute mean, std, median, p25, p75 across all companies in sector")
bullet("Z-Score Normalization: z = (raw_value − sector_mean) / sector_std → stored as SectorNormalizedFeature")
bullet("Percentile Rank: each company's position within its sector for each metric")
bullet("Adaptive Weights: historical return correlations per category (profitability, liquidity, leverage, cash flow) used to adjust rule-based weights proportionally; falls back to base weights if insufficient data")

page_break()

# ══════════════════════════════════════════════════════════════
#  SECTION 5 — SCREENSHOTS GUIDE
# ══════════════════════════════════════════════════════════════
heading1("5. Screenshots — Capture Guide")

body("Start the application with: docker-compose up --build", italic=False)
body("Frontend → http://localhost:3000  |  Swagger API → http://localhost:8000/docs", italic=True, color=MID_BLUE)
doc.add_paragraph()

make_table(
    ["Page / View", "URL", "What to Capture"],
    [
        ["Login Page",             "/login",                   "Login form with email & password fields, project branding"],
        ["Dashboard",              "/dashboard",               "Recent score-runs, quick action links, summary cards"],
        ["Company / Stock Page",   "/companies/:id",           "Financial statements, 12 computed metrics, period-over-period transitions, sector z-scores"],
        ["Search / Browse",        "/companies",               "Ticker/sector filter, company list with pagination"],
        ["Score Result",           "/score-runs/:id",          "Total score (0–100), success probability, 3-level explainability, metric contribution table"],
        ["Compare Page",           "/compare",                 "Multi-stock comparison table, 5 model scores + ensemble, ranking"],
        ["Reports / Export",       "/reports",                 "Export history list, CSV/JSON/PDF download buttons"],
        ["Forecasting",            "/forecasting",             "Upload winners, train model button, ranked prediction output"],
        ["Forecasting Detail",     "/forecasting/detail",      "Trend series charts, heatmap, parameter rankings"],
        ["Data Health",            "/data-health",             "Ingestion job list, per-row issue log (missing/outlier/duplicate)"],
        ["Validation Lab",         "/validation",              "Accuracy/F1/ROC-AUC table, confusion matrix, SHAP feature importance chart"],
        ["Labeling Lab",           "/labeling",                "Label definition list, sector benchmark config, distribution preview"],
        ["Admin Panel",            "/admin",                   "Scoring model registry, audit log table, user management"],
        ["FastAPI Swagger",        ":8000/docs",               "Full endpoint list with try-it-out forms, schema definitions"],
        ["Charts / Graphs",        "CompanyPage or ScoreResult","Recharts visualizations of financial metrics over time"],
    ],
    col_widths=[1.8, 1.7, 2.7]
)

page_break()

# ══════════════════════════════════════════════════════════════
#  SECTION 6 — GITHUB & TECH STACK
# ══════════════════════════════════════════════════════════════
heading1("6. GitHub Repository & Technologies")

heading2("6.1 Repository")
body("GitHub: salih04/capstone-financeiq")
body("Branch (active development): claude/progress-report-documentation-S054P")

heading2("6.2 Full Technology Stack")
make_table(
    ["Layer", "Technology", "Version"],
    [
        # Frontend
        ["Frontend Framework",    "React",                        "18.3.1"],
        ["Build Tool",            "Vite",                         "5.2.11"],
        ["Routing",               "React Router",                 "6.23.1"],
        ["HTTP Client",           "axios",                        "1.6.8"],
        ["Charts",                "Recharts",                     "2.12.7"],
        ["Icons",                 "lucide-react",                 "0.400.0"],
        ["E2E Testing",           "Playwright",                   "1.52.0"],
        # Backend
        ["Backend Framework",     "FastAPI",                      "0.111.0"],
        ["ASGI Server",           "Uvicorn",                      "0.29.0"],
        ["ORM",                   "SQLAlchemy",                   "2.0.30"],
        ["Migrations",            "Alembic",                      "1.13.1"],
        ["Database",              "PostgreSQL",                   "16"],
        ["DB Driver",             "psycopg2-binary",              "2.9.9"],
        ["Auth — Hashing",        "passlib + bcrypt",             "1.7.4 / 4.0.1"],
        ["Auth — Tokens",         "python-jose (HS256 JWT)",      "3.3.0"],
        ["Validation",            "Pydantic v2",                  "2.7.1"],
        # ML
        ["ML — Core",             "scikit-learn",                 "≥ 1.4.0"],
        ["ML — Boosting",         "XGBoost",                      "latest"],
        ["ML — Explainability",   "SHAP",                         "≥ 0.46.0"],
        ["Data Processing",       "pandas",                       "2.2.2"],
        ["Numerical",             "numpy",                        "≥ 1.26.0"],
        # Export & Infra
        ["PDF Generation",        "ReportLab",                    "≥ 4.1.0"],
        ["Excel I/O",             "openpyxl",                     "3.1.5"],
        ["HTML Scraping",         "BeautifulSoup4",               "4.12.3"],
        ["Async HTTP",            "httpx",                        "0.27.0"],
        ["Containerization",      "Docker + Docker Compose",      "—"],
        ["Web Server",            "Nginx (Alpine)",               "—"],
        ["Runtime",               "Python 3.12 / Node 20",        "—"],
    ],
    col_widths=[2.2, 2.4, 1.6]
)

heading2("6.3 Project Directory Structure")
code_block(
    "/home/user/capstone-financeIQ/\n"
    "├── 1.frontend/                   # React + Vite SPA\n"
    "│   ├── src/\n"
    "│   │   ├── pages/                # 16+ page components\n"
    "│   │   ├── components/           # Navbar, Sidebar, AppShell, ProtectedRoute\n"
    "│   │   ├── context/              # AuthContext (JWT state)\n"
    "│   │   ├── api/                  # axios instance with JWT interceptors\n"
    "│   │   └── App.jsx               # 18+ route definitions\n"
    "│   ├── Dockerfile                # Multi-stage: Node → Nginx\n"
    "│   ├── vite.config.js            # /api proxy to backend\n"
    "│   └── nginx.conf                # Production reverse proxy\n"
    "│\n"
    "├── 2.backend/                    # FastAPI + PostgreSQL\n"
    "│   ├── app/\n"
    "│   │   ├── main.py               # App init, routers, CORS\n"
    "│   │   ├── config.py             # Settings (DATABASE_URL, SECRET_KEY)\n"
    "│   │   ├── database.py           # SQLAlchemy engine & session\n"
    "│   │   ├── core/                 # security.py, dependencies.py\n"
    "│   │   ├── models/               # 12+ SQLAlchemy ORM models\n"
    "│   │   ├── routers/              # 11 FastAPI routers\n"
    "│   │   ├── services/             # 17 business logic services\n"
    "│   │   └── schemas/              # Pydantic request/response schemas\n"
    "│   ├── alembic/versions/         # 5 migration files (2026-04-06)\n"
    "│   ├── scripts/                  # ETL, seeding, retraining scripts\n"
    "│   ├── airflow/dags/             # forecasting_retrain_dag.py\n"
    "│   ├── tests/                    # Playwright + API contract tests\n"
    "│   ├── seed.py                   # DB seeding on startup\n"
    "│   ├── Dockerfile                # python:3.12-slim\n"
    "│   └── requirements.txt\n"
    "│\n"
    "├── 3.Datasets/                   # Historical BIST data (2020–2025)\n"
    "│   └── 2020stocks.xlsx … 2025stocks.xlsx\n"
    "│\n"
    "├── docker-compose.yml            # 3-service orchestration\n"
    "└── README.md"
)

page_break()

# ══════════════════════════════════════════════════════════════
#  APPENDIX A — ALL API ENDPOINTS
# ══════════════════════════════════════════════════════════════
heading1("Appendix A — Complete API Endpoint Reference")

heading2("Auth & Users")
make_table(
    ["Method", "Endpoint", "Description"],
    [
        ["POST", "/auth/register",          "Create new user account"],
        ["POST", "/auth/login",             "Authenticate; returns JWT access token"],
        ["GET",  "/auth/me",                "Return current authenticated user"],
        ["GET",  "/users/me",               "Current user profile details"],
        ["PUT",  "/users/me/profile",       "Update onboarding profile (user_type, risk_level, …)"],
        ["GET",  "/users/me/score-runs",    "List current user's recent score runs"],
    ],
    col_widths=[0.8, 2.5, 2.9]
)

heading2("Companies")
make_table(
    ["Method", "Endpoint", "Description"],
    [
        ["GET", "/companies",                         "Search companies (ticker, name, sector, dataset filter)"],
        ["GET", "/companies/{id}",                    "Company detail"],
        ["GET", "/companies/{id}/financials",         "Raw FinancialStatement rows"],
        ["GET", "/companies/{id}/metrics",            "12 ComputedMetric rows per period"],
        ["GET", "/companies/{id}/transitions",        "Period-over-period metric changes"],
        ["GET", "/companies/{id}/sector-scores",      "Sector z-scores and percentile ranks"],
    ],
    col_widths=[0.8, 2.8, 2.6]
)

heading2("Scoring")
make_table(
    ["Method", "Endpoint", "Description"],
    [
        ["POST", "/companies/{id}/score",     "Run scoring (rule_based | logistic | ensemble); 3-level explainability"],
        ["POST", "/scoring/common-periods",   "Intersection of available periods for multi-stock comparison"],
        ["POST", "/scoring/compare",          "Compare multiple stocks across 5 models + ensemble"],
        ["GET",  "/scoring/adaptive-weights", "Preview data-driven weight adjustments"],
        ["GET",  "/score-runs/{id}",          "Retrieve specific score-run result"],
    ],
    col_widths=[0.8, 2.5, 2.9]
)

heading2("Financials & Ingestion")
make_table(
    ["Method", "Endpoint", "Description"],
    [
        ["POST", "/financials/import-csv",     "Upload CSV; auto-calculates 12 ratios"],
        ["POST", "/ingestion/import/csv",      "Full V3 pipeline with job tracking"],
        ["GET",  "/ingestion/jobs",            "List ingestion job history"],
        ["GET",  "/ingestion/jobs/{id}",       "Single job detail + data-quality issues"],
        ["GET",  "/ingestion/dashboard",       "Data health summary"],
    ],
    col_widths=[0.8, 2.5, 2.9]
)

heading2("Forecasting")
make_table(
    ["Method", "Endpoint", "Description"],
    [
        ["POST", "/upload-data",         "Import yearly winner Excel preset (2020–2025)"],
        ["POST", "/train-model",         "Train sector success model for given year/sector"],
        ["GET",  "/get-stocks",          "Generate stock rankings via forecasting model"],
        ["GET",  "/get-parameters",      "List ranked parameters for sector"],
        ["POST", "/predict/evaluate",    "Run rolling time-CV evaluation"],
    ],
    col_widths=[0.8, 2.0, 3.4]
)

heading2("Reports, Validation, Labeling & Admin")
make_table(
    ["Method", "Endpoint", "Description"],
    [
        ["GET",    "/reports/score-runs/{id}/export.csv",           "CSV export with full breakdown"],
        ["GET",    "/reports/score-runs/{id}/export.json",          "JSON export with AI commentary"],
        ["GET",    "/reports/score-runs/{id}/export.pdf",           "Styled PDF export (ReportLab)"],
        ["POST",   "/validation/run",                               "Time-split validation on scoring model"],
        ["GET",    "/validation/models/{id}/history",               "Validation run history"],
        ["GET",    "/validation/models/{id}/feature-importances",   "SHAP / coefficient rankings"],
        ["GET",    "/labeling/definitions",                         "List label definitions"],
        ["POST",   "/labeling/definitions",                         "Create label definition"],
        ["POST",   "/labeling/definitions/{id}/preview",            "Preview label distribution"],
        ["POST",   "/labeling/definitions/{id}/activate",           "Activate label for training"],
        ["GET",    "/admin/scoring-models",                         "List all scoring models"],
        ["POST",   "/admin/scoring-models",                         "Create scoring model"],
        ["PATCH",  "/admin/scoring-models/{id}",                    "Update model"],
        ["DELETE", "/admin/scoring-models/{id}",                    "Archive model"],
        ["GET",    "/admin/audit-logs",                             "Immutable audit trail"],
        ["GET",    "/admin/users",                                  "User list"],
        ["PATCH",  "/admin/users/{id}/role",                        "Promote user to admin/analyst"],
        ["GET",    "/news/updates",                                  "Finnhub market news (12 h cache)"],
        ["GET",    "/health",                                        "API health check + version"],
    ],
    col_widths=[0.7, 2.7, 2.8]
)

page_break()

# ══════════════════════════════════════════════════════════════
#  APPENDIX B — DATABASE SCHEMA SUMMARY
# ══════════════════════════════════════════════════════════════
heading1("Appendix B — Database Schema Summary (25+ Tables)")

groups = [
    ("Core", [
        ("users",                   "Accounts · email · password_hash · role · lockout fields · onboarding profile"),
        ("companies",               "Stock master · ticker · company_name · sector_code · is_active"),
        ("financial_statements",    "Raw balance sheet + income + cash flow per (company, period)"),
        ("computed_metrics",        "12 derived ratios per (company, period) · ROA/ROE/margins/liquidity/leverage/CF"),
        ("stock_returns",           "Annual & rolling returns from .xlsx presets · price · market_cap"),
    ]),
    ("Analytics & Scoring", [
        ("metric_transitions",          "Period-over-period delta per metric (abs_change, pct_change)"),
        ("sector_benchmarks",           "Sector-wide statistics per (sector, period, feature): mean, std, percentiles"),
        ("sector_normalized_features",  "Per-company z-score & percentile rank vs sector benchmark"),
        ("score_runs",                  "Completed scoring executions · total_score · success_probability · 3-level explanation JSON"),
        ("score_details",               "Per-metric contribution breakdown per score run (L1/L2/L3 explainability)"),
    ]),
    ("Model Governance", [
        ("scoring_models",              "Model registry · model_type · version · status · feature_set · label_strategy"),
        ("scoring_model_metrics",       "Per-model feature weights · threshold_min/max · direction"),
        ("model_validation_runs",       "Time-split validation results · accuracy/precision/recall/F1/ROC-AUC"),
        ("model_feature_importances",   "SHAP / coefficient rankings per model"),
        ("label_definitions",           "Label config · sector_benchmark_type · horizon_months · threshold_rule"),
        ("audit_logs",                  "Immutable action history · actor · action_type · old/new value JSON"),
    ]),
    ("Data Pipeline", [
        ("ingestion_jobs",       "Pipeline job tracking · status · items_total/success/failed · timing"),
        ("data_quality_issues",  "Per-row anomalies · issue_type · severity · detected_at"),
    ]),
    ("Forecasting", [
        ("winner_cohort_rows",        "Yearly BIST winner stocks from .xlsx presets · returns (1w/1m/…/5y)"),
        ("sector_parameter_rankings", "Top predictive parameters per (sector, year) with scores & ranks"),
        ("forecast_runs",             "Forecasting execution record · year · sector · model_version"),
        ("forecast_predictions",      "Per-stock prediction output · score · confidence · rank · explanation"),
        ("forecast_evaluation_runs",  "Rolling window CV results · mean_rank_stability · mean_overlap_at_k"),
        ("forecast_evaluation_folds", "Per-fold metrics · train/test year ranges · rank_stability · overlap_at_k"),
        ("quarterly_fundamentals",    "Fine-grained financials per (stock, quarter) · 30+ fields"),
        ("reports",                   "Export history · user_id · score_run_id · report_type · created_at"),
    ]),
]

for group_name, tables in groups:
    heading2(group_name)
    make_table(
        ["Table", "Key Fields / Purpose"],
        [[t, d] for t, d in tables],
        col_widths=[2.4, 3.8]
    )

# ── Save ──────────────────────────────────────────────────────
out_path = "/home/user/capstone-financeIQ/FinanceIQ_Progress_Report.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
